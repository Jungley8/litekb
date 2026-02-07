"""
文档处理服务
"""

import os
import uuid
from pathlib import Path
from typing import List, Optional, Dict, Any
from loguru import logger

from app.config import settings
from app.data_models import Document, DocumentChunk


class DocumentProcessor:
    """文档处理器"""

    def __init__(self):
        self.chunk_size = settings.chunk_size
        self.chunk_overlap = settings.chunk_overlap

    async def process_document(
        self,
        file_content: bytes,
        filename: str,
        metadata: Optional[Dict[str, Any]] = None,
    ) -> Dict[str, Any]:
        """处理文档"""
        text = await self.extract_text(file_content, filename)
        chunks = self.split_chunks(text)

        return {
            "title": Path(filename).stem,
            "file_type": self.get_file_type(filename),
            "content": text[:1000],
            "chunks": chunks,
            "metadata": metadata or {},
        }

    async def extract_text(self, content: bytes, filename: str) -> str:
        """从文件中提取文本"""
        file_type = self.get_file_type(filename)

        if file_type in ["txt", "md"]:
            return content.decode("utf-8")
        elif file_type == "docx":
            return await self.extract_docx(content)
        elif file_type == "pdf":
            return await self.extract_pdf(content)
        else:
            raise ValueError(f"不支持的文件类型: {file_type}")

    def get_file_type(self, filename: str) -> str:
        """获取文件类型"""
        ext = Path(filename).suffix.lower()
        type_map = {
            ".txt": "txt",
            ".md": "md",
            ".markdown": "md",
            ".docx": "docx",
            ".pdf": "pdf",
        }
        return type_map.get(ext, "txt")

    async def extract_docx(self, content: bytes) -> str:
        """提取 DOCX 文本"""
        try:
            import docx
            from io import BytesIO

            doc = docx.Document(BytesIO(content))
            text_parts = [para.text for para in doc.paragraphs if para.text.strip()]
            return "\n".join(text_parts)
        except ImportError:
            logger.warning("python-docx 未安装")
            return ""

    async def extract_pdf(self, content: bytes) -> str:
        """提取 PDF 文本"""
        logger.warning("PDF 提取功能待实现")
        return ""

    def split_chunks(self, text: str) -> List[str]:
        """将文本分块"""
        if not text:
            return []

        chunks = []
        start = 0
        text = text.replace("\n\n", "\n")

        while start < len(text):
            end = start + self.chunk_size

            if end < len(text):
                for sep in [". ", "? ", "! ", "\n"]:
                    last_sep = text.rfind(sep, start, end)
                    if last_sep > start + self.chunk_size // 2:
                        end = last_sep + len(sep)
                        break

            chunks.append(text[start:end].strip())
            start = end - self.chunk_overlap

        return chunks


class DocumentService:
    """文档服务"""

    def __init__(self):
        self.processor = DocumentProcessor()

    def _get_session(self):
        """延迟导入 session"""
        from app.db.orm_store import orm_store
        return orm_store.get_session()

    async def create_document(
        self,
        kb_id: str,
        title: str,
        content: bytes,
        filename: str,
        metadata: Optional[Dict] = None,
    ) -> Document:
        """创建文档"""
        session = self._get_session()

        try:
            processed = await self.processor.process_document(
                content, filename, metadata
            )

            doc = Document(
                id=str(uuid.uuid4()),
                title=title,
                file_type=processed["file_type"],
                content=processed["content"],
                metadata=processed["metadata"],
                status="processing",
            )
            session.add(doc)
            session.flush()

            for i, chunk_content in enumerate(processed["chunks"]):
                chunk = DocumentChunk(
                    id=str(uuid.uuid4()),
                    doc_id=doc.id,
                    kb_id=kb_id,
                    chunk_index=i,
                    content=chunk_content,
                    metadata={"source": filename},
                )
                session.add(chunk)

            session.commit()
            doc.status = "indexed"
            session.commit()

            logger.info(f"文档 {doc.id} 创建成功")
            return doc

        except Exception as e:
            session.rollback()
            doc.status = "failed"
            doc.error_message = str(e)
            session.commit()
            raise e

        finally:
            session.close()

    def get_documents(self, kb_id: str, skip: int = 0, limit: int = 100) -> List[Document]:
        """获取知识库文档列表"""
        session = self._get_session()
        try:
            docs = (
                session.query(DocumentChunk)
                .filter(DocumentChunk.kb_id == kb_id)
                .offset(skip)
                .limit(limit)
                .all()
            )

            seen = set()
            result = []
            for doc in docs:
                if doc.doc_id not in seen:
                    seen.add(doc.doc_id)
                    full_doc = session.query(Document).get(doc.doc_id)
                    if full_doc:
                        result.append(full_doc)

            return result
        finally:
            session.close()

    def get_chunks(self, doc_id: str) -> List[DocumentChunk]:
        """获取文档分块"""
        session = self._get_session()
        try:
            chunks = (
                session.query(DocumentChunk)
                .filter(DocumentChunk.doc_id == doc_id)
                .order_by(DocumentChunk.chunk_index)
                .all()
            )
            return chunks
        finally:
            session.close()

    def delete_document(self, doc_id: str) -> bool:
        """删除文档"""
        session = self._get_session()
        try:
            session.query(DocumentChunk).filter(DocumentChunk.doc_id == doc_id).delete()
            session.query(Document).filter(Document.id == doc_id).delete()
            session.commit()
            return True
        except Exception as e:
            session.rollback()
            raise e
        finally:
            session.close()


document_service = DocumentService()
